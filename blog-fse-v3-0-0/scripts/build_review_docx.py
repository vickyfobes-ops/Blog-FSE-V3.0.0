#!/usr/bin/env python3
"""Build polished English and Chinese Word review drafts from Markdown."""

from __future__ import annotations

import argparse
import json
import re
from datetime import date
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "111111"
DARK_BLUE = "111111"
INK = "111111"
MUTED = "6A7178"
BRONZE = "333333"
LIGHT_FILL = "F4F6F9"
RULE = "D8DEE5"
TABLE_WIDTH_DXA = 9360


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, name: str, size: float | None = None, color: str | None = None,
             bold: bool | None = None, italic: bool | None = None, east_asia: str | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:cs"), name)
    fonts.set(qn("w:eastAsia"), east_asia or name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str, font_name: str, east_asia: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), east_asia)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([fonts, color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([rpr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str, font_name: str, east_asia: str) -> None:
    token = re.compile(
        r"\[([^\]]+)\]\((https?://[^\)]+)\)|\*\*([^*]+)\*\*|(?<!\*)\*([^*]+)\*(?!\*)|`([^`]+)`"
    )
    cursor = 0
    for match in token.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_font(run, font_name, east_asia=east_asia)
        link_text, url, strong, emphasis, code = match.groups()
        if link_text is not None:
            add_hyperlink(paragraph, link_text, url, font_name, east_asia)
        else:
            value = strong or emphasis or code or ""
            run = paragraph.add_run(value)
            set_font(
                run,
                "Courier New" if code else font_name,
                bold=strong is not None,
                italic=emphasis is not None,
                east_asia=east_asia,
            )
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run, font_name, east_asia=east_asia)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, "Calibri", 9, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run_el = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = "1"
    run_el.append(text_el)
    field.append(run_el)
    paragraph._p.append(field)


def add_bullet_numbering(document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multilevel = OxmlElement("w:multiLevelType")
    multilevel.set(qn("w:val"), "singleLevel")
    abstract.append(multilevel)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "279")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend([tabs, ind, spacing])
    rpr = OxmlElement("w:rPr")
    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Calibri")
    font.set(qn("w:hAnsi"), "Calibri")
    rpr.append(font)
    level.extend([start, fmt, lvl_text, jc, ppr, rpr])
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    ppr.append(num_pr)


def configure_styles(document, language: str) -> tuple[str, str]:
    # LibreOffice on macOS can ignore eastAsia-only font hints. Use the
    # installed Unicode family for every Chinese run so the review PDF and
    # Word document render identically instead of falling back to tofu boxes.
    font_name = "Arial Unicode MS" if language.startswith("zh") else "Calibri"
    east_asia = font_name
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    fonts = normal.element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), east_asia)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10, 0),
        "Heading 2": (13, BLUE, 12, 6, 1),
        "Heading 3": (12, DARK_BLUE, 8, 4, 2),
    }
    for name, (size, color, before, after, outline) in heading_tokens.items():
        style = styles[name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        fonts = style.element.rPr.get_or_add_rFonts()
        fonts.set(qn("w:ascii"), font_name)
        fonts.set(qn("w:hAnsi"), font_name)
        fonts.set(qn("w:eastAsia"), east_asia)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        ppr = style.element.get_or_add_pPr()
        outline_el = ppr.find(qn("w:outlineLvl"))
        if outline_el is None:
            outline_el = OxmlElement("w:outlineLvl")
            ppr.append(outline_el)
        outline_el.set(qn("w:val"), str(outline))

    caption = styles["Caption"]
    caption.font.name = font_name
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    fonts = caption.element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), east_asia)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(12)
    caption.paragraph_format.keep_together = True
    return font_name, east_asia


def configure_page(document, title: str, language: str, font_name: str, east_asia: str) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    label = "FSE SHOPIFY BLOG · 中文审阅稿" if language.startswith("zh") else "FSE SHOPIFY BLOG · REVIEW DRAFT"
    run = header_p.add_run(label)
    set_font(run, font_name, 8.5, MUTED, bold=True, east_asia=east_asia)
    ppr = header_p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), RULE)
    p_bdr.append(bottom)
    ppr.append(p_bdr)

    footer_p = section.footer.paragraphs[0]
    add_page_field(footer_p)


def add_review_front_matter(document, title: str, meta: dict, language: str,
                            font_name: str, east_asia: str, word_count: int) -> None:
    zh = language.startswith("zh")
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(24)
    kicker.paragraph_format.space_after = Pt(12)
    run = kicker.add_run("第一步 · 文章审阅" if zh else "STEP 1 · EDITORIAL REVIEW")
    set_font(run, font_name, 9.5, BRONZE, bold=True, east_asia=east_asia)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    set_font(run, font_name, 26, INK, bold=True, east_asia=east_asia)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("中文同步审阅稿 · 尚未上传 Shopify" if zh else "English review copy · Not uploaded to Shopify")
    set_font(run, font_name, 11, MUTED, italic=True, east_asia=east_asia)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(0)
    note.paragraph_format.space_after = Pt(16)
    note.paragraph_format.left_indent = Inches(0.18)
    note.paragraph_format.right_indent = Inches(0.18)
    text = (
        "审阅说明：本文正文、6 张配图、来源与 FAQ 已完成。你确认后，第三步才会加入已验证商品内链并生成 Shopify 安全格式。"
        if zh else
        "Review status: article, six images, sources, and FAQ are complete. Verified product links and Shopify-safe formatting are added only in Action 3 after approval."
    )
    run = note.add_run(text)
    set_font(run, font_name, 10, DARK_BLUE, bold=True, east_asia=east_asia)
    ppr = note._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), BRONZE)
    p_bdr.append(left)
    ppr.append(p_bdr)

    seo = meta.get("seo", {})
    rows = (
        [
            ("主要关键词", meta.get("primaryKeyword", "")),
            ("SEO 标题", seo.get("title", "")),
            ("URL 路径", seo.get("handle", "")),
            ("英文篇幅", f"{word_count:,} 词"),
        ]
        if zh else
        [
            ("Primary keyword", meta.get("primaryKeyword", "")),
            ("SEO title", seo.get("title", "")),
            ("URL handle", seo.get("handle", "")),
            ("English length", f"{word_count:,} words"),
        ]
    )
    table = document.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2700, 6660], indent_dxa=120)
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_FILL)
        for cell, content, bold, color in (
            (row.cells[0], label, True, DARK_BLUE),
            (row.cells[1], value, False, "333333"),
        ):
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(content)
            set_font(run, font_name, 9.5, color, bold=bold, east_asia=east_asia)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def markdown_word_count(markdown: str) -> int:
    text = re.sub(r"<!--.*?-->", "", markdown, flags=re.S)
    text = re.sub(r"!\[[^\]]*\]\([^\)]+\)", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"[#*_>`-]", " ", text)
    return len(re.findall(r"\b[\w’'-]+\b", text))


def build(markdown_path: Path, meta_path: Path, output_path: Path, language: str) -> Path:
    markdown = markdown_path.read_text(encoding="utf-8")
    meta = json.loads(meta_path.read_text(encoding="utf-8"))
    lines = re.sub(r"<!--.*?-->", "", markdown, flags=re.S).splitlines()
    title_match = re.search(r"^#\s+(.+)$", markdown, flags=re.M)
    if not title_match:
        raise ValueError("Markdown source has no H1 title")
    title = title_match.group(1).strip()
    english_path = markdown_path if language.startswith("en") else markdown_path.with_name(markdown_path.name.replace(".zh-CN", ""))
    word_count = markdown_word_count(english_path.read_text(encoding="utf-8"))

    document = Document()
    font_name, east_asia = configure_styles(document, language)
    configure_page(document, title, language, font_name, east_asia)
    add_review_front_matter(document, title, meta, language, font_name, east_asia, word_count)
    bullet_num_id = add_bullet_numbering(document)

    paragraph_lines: list[str] = []

    def flush() -> None:
        if not paragraph_lines:
            return
        text = " ".join(part.strip() for part in paragraph_lines).strip()
        paragraph_lines.clear()
        if not text:
            return
        p = document.add_paragraph()
        add_inline(p, text, font_name, east_asia)

    index = 0
    skipped_title = False
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            flush()
            index += 1
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        image = re.match(r"^!\[([^\]]*)\]\(([^\)]+)\)$", stripped)
        bullet = re.match(r"^[-+*]\s+(.+)$", stripped)
        caption = re.match(r"^\*(.+)\*$", stripped)
        if heading:
            flush()
            level = len(heading.group(1))
            if level == 1 and not skipped_title:
                skipped_title = True
            else:
                word_level = min(max(level - 1, 1), 3)
                p = document.add_paragraph(style=f"Heading {word_level}")
                add_inline(p, heading.group(2), font_name, east_asia)
        elif image:
            flush()
            image_path = (markdown_path.parent / image.group(2)).resolve()
            if not image_path.is_file():
                raise FileNotFoundError(f"Missing image: {image_path}")
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
            run = p.add_run()
            inline_shape = run.add_picture(str(image_path), width=Inches(6.0))
            # Preserve the Markdown alt text in the DOCX drawing metadata so
            # Word's accessibility tools and screen readers can describe it.
            alt_text = image.group(1).strip() or image_path.stem.replace("-", " ")
            inline_shape._inline.docPr.set("descr", alt_text)
            inline_shape._inline.docPr.set("title", alt_text)
        elif caption:
            flush()
            p = document.add_paragraph(style="Caption")
            run = p.add_run(caption.group(1))
            set_font(run, font_name, 9, MUTED, italic=True, east_asia=east_asia)
        elif bullet:
            flush()
            p = document.add_paragraph()
            apply_numbering(p, bullet_num_id)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            add_inline(p, bullet.group(1), font_name, east_asia)
        else:
            paragraph_lines.append(stripped)
        index += 1
    flush()

    document.core_properties.title = title
    document.core_properties.subject = "FSE Shopify blog review draft"
    document.core_properties.author = "FSE Shopify Blog"
    document.core_properties.comments = "Review draft; not uploaded to Shopify."
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-dir", required=True, type=Path)
    args = parser.parse_args()
    root = args.source_dir.expanduser().resolve()
    meta = next(root.glob("*.meta.json"))
    english = next(path for path in root.glob("*.md") if not path.name.endswith((".zh-CN.md", ".review.md")))
    chinese = next(root.glob("*.zh-CN.md"))
    slug = english.stem
    outputs = [
        build(english, meta, root / f"{slug}.en.review.docx", "en-US"),
        build(chinese, meta, root / f"{slug}.zh-CN.review.docx", "zh-CN"),
    ]
    print("\n".join(str(path) for path in outputs))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
